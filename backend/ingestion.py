"""
Multi-modal Document Processor
Extracts text, tables (as Markdown), and images (via OCR) from documents.
Supports: PDFs, scanned PDFs, tables in images, Excel spreadsheets, DOCX, TXT.
"""
import os
import json
import io
from typing import List, Dict, Any, Optional
from datetime import datetime

# Document parsing
import fitz  # PyMuPDF - pip install PyMuPDF
import docx
import openpyxl  # Excel support

# Table handling  
import pandas as pd

# OCR for images
from rapidocr_onnxruntime import RapidOCR

# Image preprocessing and table detection
import cv2
import numpy as np
from img2table.document import Image as Img2TableImage
from img2table.ocr import TesseractOCR

from fastapi import UploadFile


class DocumentProcessor:
    """
    Multi-modal document processor supporting:
    - PDF: Text, Tables (as Markdown), Images (via OCR)
    - Scanned PDFs: Full-page OCR when text extraction fails
    - Tables in images: Structured table detection
    - Low-quality scans: Image preprocessing for better OCR
    - DOCX: Text extraction
    - TXT: Plain text
    """
    
    def __init__(self, upload_dir: str = "uploads", models_dir: str = "models"):
        self.upload_dir = upload_dir
        self.models_dir = models_dir
        self.progress_file = os.path.join(models_dir, "ingestion_status.json")
        
        os.makedirs(self.upload_dir, exist_ok=True)
        os.makedirs(self.models_dir, exist_ok=True)
        
        # Initialize OCR engines (lazy load on first use)
        self._ocr_engine = None
        self._img2table_ocr = None
    
    @property
    def ocr_engine(self):
        """Lazy-load OCR engine for general text extraction."""
        if self._ocr_engine is None:
            print("Initializing RapidOCR engine...")
            self._ocr_engine = RapidOCR()
        return self._ocr_engine
    
    @property
    def img2table_ocr(self):
        """Lazy-load img2table OCR for table structure detection."""
        if self._img2table_ocr is None:
            print("Initializing TesseractOCR for table detection...")
            self._img2table_ocr = TesseractOCR(lang="eng")
        return self._img2table_ocr
    
    def preprocess_image(self, img_bytes: bytes) -> bytes:
        """
        Preprocess image for better OCR results on low-quality scans.
        Steps: Grayscale → Denoise → Adaptive threshold (binarization)
        """
        try:
            nparr = np.frombuffer(img_bytes, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_GRAYSCALE)
            
            if img is None:
                return img_bytes
            
            # Denoise
            img = cv2.fastNlMeansDenoising(img, h=10)
            
            # Adaptive threshold for binarization (helps with uneven lighting)
            img = cv2.adaptiveThreshold(
                img, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                cv2.THRESH_BINARY, 11, 2
            )
            
            # Encode back to PNG bytes
            _, encoded = cv2.imencode('.png', img)
            return encoded.tobytes()
        except Exception as e:
            print(f"  [Preprocess Warning] Image preprocessing failed: {e}")
            return img_bytes
    
    def extract_tables_from_image(self, img_bytes: bytes) -> List[str]:
        """
        Detect and extract tables from an image using img2table.
        Returns a list of Markdown-formatted tables.
        """
        tables_md = []
        try:
            # Save to temp file as img2table needs a file path
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                tmp.write(img_bytes)
                tmp_path = tmp.name
            
            # Detect tables
            img_doc = Img2TableImage(src=tmp_path)
            extracted = img_doc.extract_tables(
                ocr=self.img2table_ocr,
                borderless_tables=True,  # Also detect borderless tables
                min_confidence=50
            )
            
            # Clean up temp file
            os.unlink(tmp_path)
            
            # Convert to Markdown
            for table in extracted:
                if table.df is not None and not table.df.empty:
                    md = table.df.to_markdown(index=False)
                    tables_md.append(md)
                    print(f"  [img2table] Extracted table with {len(table.df)} rows")
                    
        except Exception as e:
            print(f"  [img2table Warning] Table extraction failed: {e}")
        
        return tables_md

    async def save_file(self, file: UploadFile) -> str:
        """Save uploaded file to disk."""
        file_path = os.path.join(self.upload_dir, file.filename)
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        return file_path

    def update_progress(self, status: str, current: int, total: int, message: str):
        """
        Write ingestion progress to a JSON file for frontend polling.
        
        Args:
            status: "processing", "complete", or "error"
            current: Current block number being processed
            total: Total number of blocks
            message: Human-readable status message
        """
        progress_data = {
            "status": status,
            "current_block": current,
            "total_blocks": total,
            "message": message,
            "percent": int((current / total) * 100) if total > 0 else 0
        }
        
        with open(self.progress_file, "w") as f:
            json.dump(progress_data, f)
        
        print(f"[Progress] {message} ({current}/{total})")

    def get_progress(self) -> Dict[str, Any]:
        """Read current ingestion progress from file."""
        if os.path.exists(self.progress_file):
            with open(self.progress_file, "r") as f:
                return json.load(f)
        return {"status": "idle", "current_block": 0, "total_blocks": 0, "message": "No ingestion in progress"}

    def extract_structured_content(self, page: fitz.Page) -> List[Dict[str, Any]]:
        """
        Extract structured content from a single PDF page.
        
        Returns a list of content items, each with:
        - type: "text", "table", or "image_text"
        - content: The extracted content
        - source: Description of where this came from
        """
        content_items = []
        page_num = page.number + 1  # fitz uses 0-indexed pages
        
        # 1. TABLE EXTRACTION
        try:
            tables = page.find_tables()
            if tables and tables.tables:
                for i, table in enumerate(tables.tables):
                    # Extract table data
                    table_data = table.extract()
                    
                    if table_data and len(table_data) > 1:  # Has header + at least one row
                        # Convert to DataFrame for Markdown conversion
                        df = pd.DataFrame(table_data[1:], columns=table_data[0])
                        markdown_table = df.to_markdown(index=False)
                        
                        content_items.append({
                            "type": "table",
                            "content": markdown_table,
                            "source": f"Table {i+1} on page {page_num}",
                            "page": page_num
                        })
                        print(f"  [Table] Extracted table {i+1} from page {page_num}")
        except Exception as e:
            print(f"  [Table Warning] Could not extract tables from page {page_num}: {e}")
        
        # 2. IMAGE/OCR EXTRACTION (with table detection and preprocessing)
        try:
            images = page.get_images(full=True)
            for img_idx, img_info in enumerate(images):
                xref = img_info[0]
                
                try:
                    # Extract image bytes
                    base_image = page.parent.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Preprocess for better OCR on low-quality scans
                    processed_bytes = self.preprocess_image(image_bytes)
                    
                    # First: Try to detect tables in the image
                    tables_md = self.extract_tables_from_image(image_bytes)
                    for t_idx, table_md in enumerate(tables_md):
                        content_items.append({
                            "type": "image_table",
                            "content": table_md,
                            "source": f"Table from image {img_idx+1} on page {page_num}",
                            "page": page_num
                        })
                    
                    # Second: Run general OCR for non-table text
                    result, _ = self.ocr_engine(processed_bytes)
                    
                    if result:
                        # Combine all OCR text
                        ocr_text = " ".join([line[1] for line in result])
                        
                        # Only include if meaningful text found (>20 chars, excluding if table already extracted)
                        if len(ocr_text.strip()) > 20 and not tables_md:
                            content_items.append({
                                "type": "image_text",
                                "content": ocr_text.strip(),
                                "source": f"OCR from image {img_idx+1} on page {page_num}",
                                "page": page_num
                            })
                            print(f"  [OCR] Extracted text from image {img_idx+1} on page {page_num}")
                except Exception as img_e:
                    print(f"  [OCR Warning] Could not process image {img_idx+1} on page {page_num}: {img_e}")
        except Exception as e:
            print(f"  [OCR Warning] Could not get images from page {page_num}: {e}")
        
        # 3. NORMAL TEXT EXTRACTION
        try:
            text = page.get_text("text")
            if text and text.strip():
                content_items.append({
                    "type": "text",
                    "content": text.strip(),
                    "source": f"Text from page {page_num}",
                    "page": page_num
                })
        except Exception as e:
            print(f"  [Text Warning] Could not extract text from page {page_num}: {e}")
        
        # 4. SCANNED PDF DETECTION (Full-page OCR fallback)
        # If no text extracted but page has content, treat as scanned page
        text_content = [c for c in content_items if c["type"] == "text"]
        has_text = any(len(c["content"]) > 50 for c in text_content)
        
        if not has_text:
            try:
                print(f"  [Scanned] Page {page_num} appears to be scanned. Running full-page OCR...")
                # Render page as high-DPI image
                pix = page.get_pixmap(dpi=300)
                img_bytes = pix.tobytes("png")
                
                # Preprocess the full page image
                processed_bytes = self.preprocess_image(img_bytes)
                
                # Try table detection on full page
                tables_md = self.extract_tables_from_image(img_bytes)
                for t_idx, table_md in enumerate(tables_md):
                    content_items.append({
                        "type": "scanned_table",
                        "content": table_md,
                        "source": f"Table from scanned page {page_num}",
                        "page": page_num
                    })
                
                # Run OCR on full page
                result, _ = self.ocr_engine(processed_bytes)
                if result:
                    ocr_text = " ".join([line[1] for line in result])
                    if len(ocr_text.strip()) > 50:
                        content_items.append({
                            "type": "scanned_text",
                            "content": ocr_text.strip(),
                            "source": f"OCR from scanned page {page_num}",
                            "page": page_num
                        })
                        print(f"  [Scanned] Extracted {len(ocr_text)} chars from page {page_num}")
            except Exception as e:
                print(f"  [Scanned Warning] Full-page OCR failed on page {page_num}: {e}")
        
        return content_items

    def create_smart_chunks(
        self, 
        content_items: List[Dict[str, Any]], 
        chunk_size: int = 500, 
        overlap: int = 50
    ) -> List[Dict[str, Any]]:
        """
        Intelligently chunk content based on type.
        
        Rules:
        - Tables and OCR content: Treat as ATOMIC (no splitting)
        - Normal text: Split by chunk_size with overlap
        """
        processed_chunks = []
        
        for item in content_items:
            content_type = item.get("type", "text")
            content = item.get("content", "")
            page = item.get("page", 1)
            source = item.get("source", "unknown")
            
            if content_type in ("table", "image_text"):
                # ATOMIC CHUNK: Do not split tables or OCR content
                chunk_data = {
                    "text": content,
                    "page": page,
                    "source": source,
                    "type": content_type
                }
                if "metadata" in item:
                    chunk_data["metadata"] = item["metadata"]
                processed_chunks.append(chunk_data)
            else:
                # NORMAL TEXT: Split with overlap
                text = content
                start = 0
                chunk_idx = 0
                
                while start < len(text):
                    end = min(start + chunk_size, len(text))
                    chunk_text = text[start:end]
                    
                    chunk_data = {
                        "text": chunk_text,
                        "page": page,
                        "source": f"{source} (chunk {chunk_idx + 1})",
                        "type": "text"
                    }
                    if "metadata" in item:
                        chunk_data["metadata"] = item["metadata"]
                    processed_chunks.append(chunk_data)
                    
                    start += (chunk_size - overlap)
                    chunk_idx += 1
        
        return processed_chunks

    def extract_text(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Multi-modal text extraction from file.
        
        For PDFs: Extracts text, tables, and OCR from images.
        For DOCX/TXT: Extracts plain text.
        """
        ext = os.path.splitext(file_path)[1].lower()
        all_content = []

        if ext == ".pdf":
            print(f"[PDF] Processing: {file_path}")
            doc = fitz.open(file_path)
            total_pages = len(doc)
            
            self.update_progress("processing", 0, total_pages, f"Starting PDF extraction ({total_pages} pages)...")
            
            for page_idx, page in enumerate(doc):
                print(f"[PDF] Processing page {page_idx + 1}/{total_pages}")
                page_content = self.extract_structured_content(page)
                all_content.extend(page_content)
                
                # Update progress every 5 pages
                if (page_idx + 1) % 5 == 0 or page_idx == total_pages - 1:
                    self.update_progress(
                        "processing", 
                        page_idx + 1, 
                        total_pages, 
                        f"Extracted page {page_idx + 1}/{total_pages}"
                    )
            
            doc.close()
            self.update_progress("complete", total_pages, total_pages, "PDF extraction complete")
            
        elif ext == ".docx":
            print(f"[DOCX] Processing: {file_path}")
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            all_content.append({
                "type": "text",
                "content": "\n".join(full_text),
                "page": 1,
                "source": "DOCX document"
            })
            
        elif ext == ".txt":
            print(f"[TXT] Processing: {file_path}")
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            
            all_content.append({
                "type": "text",
                "content": text,
                "page": 1,
                "source": "TXT file"
            })
        
        elif ext in [".xlsx", ".xls"]:
            print(f"[EXCEL] Processing: {file_path}")
            all_content.extend(self.extract_excel_content(file_path))
        
        elif ext == ".csv":
            print(f"[CSV] Processing: {file_path}")
            all_content.extend(self.extract_csv_content(file_path))
        
        return all_content

    def extract_excel_content(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract content from Excel files, sheet by sheet.
        Capture file modification time for version control.
        """
        content_items = []
        try:
            # Get file stats for version control
            stats = os.stat(file_path)
            mod_time = datetime.fromtimestamp(stats.st_mtime).isoformat()
            
            # Load workbook
            wb = openpyxl.load_workbook(file_path, data_only=True)
            
            for sheet_name in wb.sheetnames:
                # Read sheet into DataFrame
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                if not df.empty:
                    # Convert to Markdown
                    md_table = df.to_markdown(index=False)
                    
                    content_items.append({
                        "type": "table",
                        "content": md_table,
                        "source": f"Excel Sheet: {sheet_name}",
                        "page": 1,
                        "metadata": {
                            "sheet_name": sheet_name,
                            "last_modified": mod_time,
                            "file_type": "excel"
                        }
                    })
                    print(f"  [Excel] Extracted sheet '{sheet_name}' with {len(df)} rows")
            
            wb.close()
            
        except Exception as e:
            print(f"  [Excel Warning] Failed to process {file_path}: {e}")
            
        return content_items

    def extract_csv_content(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract content from CSV files.
        Capture file modification time for version control.
        """
        content_items = []
        try:
            # Get file stats
            stats = os.stat(file_path)
            mod_time = datetime.fromtimestamp(stats.st_mtime).isoformat()
            
            df = pd.read_csv(file_path)
            
            if not df.empty:
                md_table = df.to_markdown(index=False)
                
                content_items.append({
                    "type": "table",
                    "content": md_table,
                    "source": "CSV File",
                    "page": 1,
                    "metadata": {
                        "last_modified": mod_time,
                        "file_type": "csv"
                    }
                })
                print(f"  [CSV] Extracted {len(df)} rows")
                
        except Exception as e:
            print(f"  [CSV Warning] Failed to process {file_path}: {e}")
            
        return content_items

    def chunk_text(self, raw_content: List[Dict[str, Any]], chunk_size: int = 500, overlap: int = 50) -> List[Dict[str, Any]]:
        """
        Wrapper for smart chunking. Maintains backward compatibility.
        """
        return self.create_smart_chunks(raw_content, chunk_size, overlap)
